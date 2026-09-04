"""Create final Systems-ready experiment tables from frozen CSV evidence."""
from pathlib import Path
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'figures_tables'
EXP = ROOT / 'results_frozen' / 'exports'
BLUE = '1F4E79'; LIGHT = 'EAF1F8'; MID = 'D9E2F3'; INK = '1F2937'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def cell_margin(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); m = tcPr.first_child_found_in('w:tcMar')
    if m is None: m = OxmlElement('w:tcMar'); tcPr.append(m)
    for side, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=m.find(qn('w:'+side))
        if node is None: node=OxmlElement('w:'+side); m.append(node)
        node.set(qn('w:w'),str(value)); node.set(qn('w:type'),'dxa')
def borders(table):
    tblPr=table._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'B7C3D0'); b.append(e)
    tblPr.append(b)
def set_repeat_row(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)
def set_cell_width(cell, width):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(width)); tcW.set(qn('w:type'),'dxa')
def make_doc():
    d=Document(); sec=d.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Inches(11); sec.page_height=Inches(8.5)
    sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.55); sec.right_margin=Inches(.55)
    normal=d.styles['Normal']; normal.font.name='Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); normal.font.size=Pt(8.5)
    return d
def add_title(d,text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(INK)
def add_note(d,text):
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(0)
    r=p.add_run('Note. '); r.bold=True; r.font.size=Pt(7.4)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(7.4)
def add_table(d, headers, rows, widths, font=7.6):
    t=d.add_table(rows=1, cols=len(headers)); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER; borders(t)
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=h; shade(c,BLUE); set_cell_width(c,widths[j]); cell_margin(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.name='Times New Roman'; r.font.size=Pt(font)
    set_repeat_row(t.rows[0])
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=str(v); set_cell_width(c,widths[j]); cell_margin(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i%2==1: shade(c,'F7F9FB')
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(font)
    return t
def read_csv(name):
    with open(EXP/name,encoding='utf-8-sig',newline='') as f:return list(csv.DictReader(f))
def f(x,n=1):return f'{float(x):.{n}f}'
def ci(mean,lo,hi,n=1):return f'{f(mean,n)} [{f(lo,n)}, {f(hi,n)}]'

def table1():
    d=make_doc(); add_title(d,'Table 1. Calibration and Experimental Contract')
    rows=[
      ['Demand timing','Sao Paulo Olist purchase timestamps; 15,540 orders','Proxy for demand timing, not last-mile release timestamps.'],
      ['Spatial demand','12 fixed calibration-only abstract zones','Zone weights are fitted on calibration data; platform origin remains an engineering assumption.'],
      ['Weather','Sao Paulo GFS 06 UTC forecasts paired with ERA5 actual 10 m wind','20-minute matched replay; forecast-minus-actual speed bias = 0.488 m/s.'],
      ['Weather gates','P50 = 0.630 m/s; P80 = 1.163 m/s; primary P90 = 1.479 m/s','P90 fixed before the 60-scenario primary evaluation; P80 is a boundary condition.'],
      ['Energy','External DJI Matrice 100 public mission-energy envelope','Transfer envelope only; no claim of same-city or target-aircraft telemetry.'],
      ['Policy factors','Planning access: always vs event-triggered; inertia: fixed vs adaptive','Five-policy evaluation also includes nonfactorial myopic reference.'],
      ['Safety','Shared ForcedSafetyTrigger validates the released first slot','Unsafe task forces safe replan or deterministic hold/swap/defer for every policy.'],
      ['Primary evaluation','60 shared Sao Paulo order-weather holdout replications; common random numbers','Scenario-level paired Student-t CI is primary; paired bootstrap and order-level GLMM are diagnostics.']]
    add_table(d,['Component','Frozen evidence or setting','Scope and interpretation'],rows,[1800,3400,5400],7.7)
    add_note(d,'Input distributions are empirically informed. VIP labels, platform movement, W_3, and H_s remain scenario or governance parameters rather than observational estimates.')
    d.save(OUT/'Table1_Calibration_and_Experimental_Contract.docx')

def table2():
    data=read_csv('primary_policy_summary.csv'); by={}
    for r in data:by.setdefault(r['Policy'],{})[r['Outcome']]=r
    d=make_doc(); add_title(d,'Table 2. Primary Sao Paulo Order-Weather Holdout Policy Evidence (60 Shared Replications)')
    labels=[('AlwaysFixed','Always fixed'),('ETFixed','ET fixed'),('AlwaysAdaptive','Always adaptive'),('ETAdaptive','ET adaptive'),('AlwaysMyopic','Myopic reference')]
    rows=[]
    for k,label in labels:
        x=by[k]; rows.append([label,ci(*map_ci(x['VIPOnTimeRate'])),ci(*map_ci(x['CompletionRate'])),ci(*map_ci(x['EnergyPerOrder'])),ci(*map_ci(x['HsPolicy'])),ci(*map_ci(x['PolicyPlanningReleaseRate'])),f(x['UnsafeExecutionRate']['Mean'])])
    add_table(d,['Policy','VIP on-time\n% [95% CI]','Completion\n% [95% CI]','Energy/order Wh\n[95% CI]','Policy H_s\n[95% CI]','Ordinary access\n% [95% CI]','Infeasible execution\nunder shared guard %'],rows,[1420,1400,1400,1450,1450,1650,650],6.6)
    c=read_csv('primary_paired_contrasts.csv'); lookup={(r['Contrast'],r['Outcome']):r for r in c}
    e=lookup[('ETFixed minus AlwaysFixed','CompletionRate')]; v=lookup[('ETFixed minus AlwaysFixed','VIPOnTimeRate')]; hp=lookup[('ETFixed minus AlwaysFixed','HsPolicy')]; ac=lookup[('ETFixed minus AlwaysFixed','PolicyPlanningReleaseRate')]
    add_note(d,f"Primary paired contrast, ET fixed minus Always fixed: completion {ci(e['MeanDifference'],e['T95Lower'],e['T95Upper'])} pp; policy H_s {ci(hp['MeanDifference'],hp['T95Lower'],hp['T95Upper'])}; ordinary access {ci(ac['MeanDifference'],ac['T95Lower'],ac['T95Upper'])} pp. VIP on-time difference is {ci(v['MeanDifference'],v['T95Lower'],v['T95Upper'])} pp under the Student-t analysis and [{f(v['Bootstrap95Lower'])}, {f(v['Bootstrap95Upper'])}] pp under paired bootstrap; its primary t interval marginally crosses zero. All policies use the shared guard and had zero infeasible execution. A 3 pp completion gain is a transparent management reference, not a calibrated economic threshold or a significance substitute: the point estimate exceeds it, but its 95% CI includes gains below 3 pp.")
    d.save(OUT/'Table2_Primary_Holdout_Policy_Evidence.docx')
def map_ci(x): return (x['Mean'],x['T95Lower'],x['T95Upper'])

def table3():
    a=read_csv('access_mechanism_boundaries.csv'); w=read_csv('inertia_sensitivity.csv'); g=read_csv('safety_guard_ablation.csv')
    d=make_doc(); add_title(d,'Table 3. Mechanism, Boundary, and Safety Evidence')
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run('Panel A. Access-rule mechanism and boundary conditions'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE)
    rows=[]
    for x in a:
        rows.append([x['Condition'],x['Scenarios'],ci(x['CompletionDelta'],x['CompletionLower'],x['CompletionUpper']),ci(x['EnergyDelta'],x['EnergyLower'],x['EnergyUpper']),ci(x['PolicyHsDelta'],x['PolicyHsLower'],x['PolicyHsUpper']),ci(x['AccessDelta'],x['AccessLower'],x['AccessUpper']),'0'])
    add_table(d,['Condition','n','Completion delta\npp [95% CI]','Energy delta Wh\n[95% CI]','Policy H_s delta\n[95% CI]','Access delta pp\n[95% CI]','Infeasible execution\nwith shared guard ON %'],rows,[2140,450,1520,1450,1540,1540,760],6.2)
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3); r=p.add_run('Panel B. Adaptive inertia sensitivity: ET adaptive minus ET fixed'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE)
    rows=[]
    for x in w: rows.append([x['Condition'],x['Scenarios'],ci(x['VIPAdaptiveMinusFixed'],x['VIPLower'],x['VIPUpper']),ci(x['CompletionAdaptiveMinusFixed'],x['CompletionLower'],x['CompletionUpper']),ci(x['PolicyHsAdaptiveMinusFixed'],x['PolicyHsLower'],x['PolicyHsUpper']),'0'])
    add_table(d,['Inertia range','n','VIP delta pp\n[95% CI]','Completion delta pp\n[95% CI]','Policy H_s delta\n[95% CI]','Infeasible execution\nwith shared guard ON %'],rows,[3060,480,1740,1740,1740,600],6.2)
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3); r=p.add_run('Panel C. Paired safety-guard ablation'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE)
    rows=[]
    for x in g:rows.append([x['Policy'],x['Scenarios'],f(x['GuardOnUnsafePct']),f(x['GuardOffUnsafePct']),ci(x['OffMinusOnUnsafePct'],x['T95Lower'],x['T95Upper']),ci(x['GuardOnReleaseP95Seconds'],x['LatencyLower'],x['LatencyUpper'],3)])
    add_table(d,['Policy','n','Guard on unsafe\n%','Guard off unsafe\n%','Off - on unsafe pp\n[95% CI]','Guard-on P95 release time s\n[95% CI]'],rows,[1800,600,1400,1400,2100,2200],7.0)
    add_note(d,'Panel A reports ET fixed minus Always fixed. Panel B does not show a stable adaptive advantage in any tested inertia range. All intervals are scenario-level paired Student-t 95% confidence intervals. Boundary and ablation analyses use 20 shared replications and are not used to reselect the primary policy.')
    d.save(OUT/'Table3_Mechanism_Boundary_and_Safety_Evidence.docx')

def table4():
    fleet=read_csv('fleet_size_boundaries.csv')
    role=read_csv('fleet_role_utilisation.csv')
    het=read_csv('input_heterogeneity_diagnostics.csv')
    d=make_doc(); add_title(d,'Table 4. Fleet-Size Boundary and Descriptive Input Heterogeneity')
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run('Panel A. Fleet-size boundary: ET fixed minus Always fixed'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE)
    rows=[]
    for x in fleet:
        rows.append([x['FleetSize'],x['Scenarios'],ci(x['VIPDelta'],x['VIPLower'],x['VIPUpper']),ci(x['CompletionDelta'],x['CompletionLower'],x['CompletionUpper']),ci(x['EnergyDelta'],x['EnergyLower'],x['EnergyUpper']),ci(x['PolicyHsDelta'],x['PolicyHsLower'],x['PolicyHsUpper']),ci(x['AccessDelta'],x['AccessLower'],x['AccessUpper']),'0'])
    add_table(d,['D','n','VIP delta pp\n[95% CI]','Completion delta pp\n[95% CI]','Energy delta Wh\n[95% CI]','Policy H_s delta\n[95% CI]','Access delta pp\n[95% CI]','Infeasible execution\nwith shared guard ON %'],rows,[470,470,1450,1580,1450,1500,1500,700],6.0)
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3); r=p.add_run('Panel B. ET-fixed role utilisation'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(BLUE)
    rows=[]
    role_names={'1':'UAV 1: capacity support','2':'UAV 2: priority service','3':'UAV 3: contingency reserve','4':'UAV 4: capacity support'}
    for x in role:
        if x['Policy']=='ETFixed': rows.append([x['FleetSize'],role_names[str(x['DroneIndex'])],ci(x['MeanActions'],x['T95Lower'],x['T95Upper'],2)])
    add_table(d,['Fleet size D','UAV role','Executed sorties per episode [95% CI]'],rows,[1600,3100,4660],7.2)
    add_note(d,'Detailed outcome-blind input heterogeneity is reported separately as Appendix Table A6 so that this main-text table retains one message: fleet-size boundary and role use.')
    d.save(OUT/'Table4_Fleet_Size_and_Input_Heterogeneity.docx')

    d=make_doc(); add_title(d,'Appendix Table A6. Outcome-Blind Descriptive Input Heterogeneity')
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(3); r=p.add_run('ET fixed minus Always fixed; not a confirmatory subgroup or interaction test'); r.italic=True; r.font.size=Pt(8.0); r.font.color.rgb=RGBColor.from_string(INK)
    rows=[]
    for x in het:
        if x['Outcome'] in ('VIP','Completion','Energy','PolicyHs','Access'):
            rows.append([x['InputCovariate'],x['InputStratum'],x['Scenarios'],f(x['CovariateMean'],3),x['Outcome'],ci(x['ETFixedMinusAlwaysFixed'],x['T95Lower'],x['T95Upper'])])
    add_table(d,['Input covariate','Stratum','n','Mean covariate','Outcome','ET fixed - Always fixed [95% CI]'],rows,[1950,1050,520,1250,1050,3540],6.4)
    add_note(d,'The heterogeneity split ranks input covariates and scenario seeds after the primary analysis and without inspecting policy outcomes. Each stratum contains 30 primary scenarios. It is descriptive, does not alter the primary inference, and does not estimate interaction or moderator effects.')
    d.save(OUT/'Appendix_Table_A6_Outcome_Blind_Input_Heterogeneity.docx')

if __name__=='__main__':
    OUT.mkdir(parents=True,exist_ok=True); table1(); table2(); table3(); table4(); print('Created final submission tables in',OUT)
